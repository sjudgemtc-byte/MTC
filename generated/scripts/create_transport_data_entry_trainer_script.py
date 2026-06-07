from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(
    "/Users/simonjudge/Documents/MTC/04_Classroom_Activities/2026-06-08_MTC_Redfern_RED-01_Pathway_Transport-Logistics-Data-Entry_Trainer-Script.docx"
)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)


def heading(doc: Document, text: str, size: int = 15) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    if not p.runs:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    else:
        p.runs[0].text = text


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    if not p.runs:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    else:
        p.runs[0].text = text


def build() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Trainer Script")
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    s = subtitle.add_run("Transport and Logistics Data Entry Activity")
    s.font.name = "Arial"
    s.font.size = Pt(12)

    meta = doc.add_table(rows=2, cols=4)
    meta.rows[0].cells[0].text = "Class"
    meta.rows[0].cells[1].text = "RED-01 Pathway"
    meta.rows[0].cells[2].text = "Trainer"
    meta.rows[0].cells[3].text = "Simon Judge"
    meta.rows[1].cells[0].text = "Topic"
    meta.rows[1].cells[1].text = "Transport and Logistics"
    meta.rows[1].cells[2].text = "Focus"
    meta.rows[1].cells[3].text = "Digital literacy and data entry"
    style_table(meta)
    for cell in (meta.rows[0].cells[0], meta.rows[0].cells[2], meta.rows[1].cells[0], meta.rows[1].cells[2]):
        shade(cell, "EDEDED")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    heading(doc, "Purpose")
    body(
        doc,
        "This activity helps learners choose a suitable application, copy transport and logistics data accurately, check times and pay rates carefully, and use AI only after they have made their own first version.",
    )

    heading(doc, "Files To Prepare")
    bullet(doc, "Print or open the student activity sheet.")
    bullet(doc, "Have the simple template ready for learners who need a lower-language version.")
    bullet(doc, "Choose whether learners will use Word, Google Docs, Excel, Google Sheets, or Numbers.")
    bullet(doc, "Have the model heading ready: Transport and Logistics Data Entry Challenge.")

    heading(doc, "Before Class")
    numbered(doc, "Open the sample file or worksheet so you can model the task.")
    numbered(doc, "Check that learners can access the app you want them to use.")
    numbered(doc, "Write the important rule on the board: Do the first copy yourself. No AI for the first version.")
    numbered(doc, "Decide which version to give each learner: full sheet or simple template.")

    heading(doc, "Simple Delivery Script")
    script = doc.add_table(rows=1, cols=3)
    script.rows[0].cells[0].text = "Stage"
    script.rows[0].cells[1].text = "Approx. time"
    script.rows[0].cells[2].text = "What to say / do"
    for cell in script.rows[0].cells:
        shade(cell, "D9EAF7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True

    stages = [
        (
            "1. Introduce the task",
            "5 min",
            "Tell learners they are practising workplace-style data entry. Explain that accuracy is more important than speed. Point out the no-AI-first rule.",
        ),
        (
            "2. Choose the application",
            "5 min",
            "Show the app choices. Help learners choose a tool that matches their confidence level. Word or Google Docs is suitable for simpler learners. Excel, Sheets, or Numbers suits learners ready for rows and columns.",
        ),
        (
            "3. Model the heading",
            "5 min",
            "Type the heading together: Transport and Logistics Data Entry Challenge. Show how to make the heading bold.",
        ),
        (
            "4. Model the first row",
            "10 min",
            "Copy the first row together. Say each part aloud: job title, suburb, time, pay, and licence. Remind learners to check capital letters, numbers, and dollar signs.",
        ),
        (
            "5. Independent copying",
            "20-30 min",
            "Learners continue copying the remaining rows. Circulate and prompt learners to compare carefully before moving to the next row.",
        ),
        (
            "6. Checking stage",
            "10 min",
            "Ask learners to check the suburb names, start and finish times, pay rates, and jobs that need a licence. Use the checklist on the sheet.",
        ),
        (
            "7. AI editing stage",
            "10-15 min",
            "Only after the first version is complete, allow learners to use the AI prompt from the sheet to check spelling, formatting, and unclear headings.",
        ),
        (
            "8. Reflection and save",
            "5-10 min",
            "Learners answer the reflection questions and save the file using a clear name.",
        ),
    ]

    for stage, time, prompt in stages:
        row = script.add_row().cells
        row[0].text = stage
        row[1].text = time
        row[2].text = prompt

    style_table(script)

    heading(doc, "Useful Teacher Prompts")
    bullet(doc, "What application did you choose, and why?")
    bullet(doc, "Check the row number before you copy.")
    bullet(doc, "Does the suburb start with a capital letter?")
    bullet(doc, "Did you keep the dollar sign and two decimal places?")
    bullet(doc, "Which jobs need a licence?")
    bullet(doc, "Have you finished your own version before asking AI to help edit it?")

    heading(doc, "Differentiation")
    bullet(doc, "Use the simple template for learners who need fewer columns and less reading.")
    bullet(doc, "Give learners the first one or two rows already started if they need extra support.")
    bullet(doc, "Pair stronger learners with weaker learners for checking, but still ask each learner to do their own first copy.")
    bullet(doc, "For extension, ask stronger learners to sort the jobs by pay rate or identify the highest pay rate after copying.")

    heading(doc, "Evidence And Observation")
    bullet(doc, "Learner chose an appropriate application.")
    bullet(doc, "Learner copied rows with varying levels of accuracy.")
    bullet(doc, "Learner checked spelling, times, pay, and licence details.")
    bullet(doc, "Learner followed the rule of completing the first version before AI editing.")
    bullet(doc, "Learner saved the file with a clear name.")

    note = doc.add_table(rows=1, cols=1)
    note.cell(0, 0).text = (
        "Audit-safe note: This activity provides formative evidence of application choice, basic data entry, checking accuracy, file saving, and supported AI editing after a learner-produced first attempt."
    )
    style_table(note)
    shade(note.cell(0, 0), "FFF2CC")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build()
