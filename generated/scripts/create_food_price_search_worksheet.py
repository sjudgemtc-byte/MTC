from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(
    "/Users/simonjudge/Documents/MTC/04_Classroom_Activities/2026-06-06_MTC_Redfern_RED-01_Pathway_Food-Price-Search_Student_Worksheet.docx"
)

FOOD_ITEMS = [
    ("Bananas", "bananas 1kg"),
    ("Apples", "apples 1kg"),
    ("Oranges", "oranges 1kg"),
    ("Washed potatoes", "washed potatoes 2kg"),
    ("Brown onions", "brown onions 1kg"),
    ("Carrots", "carrots 1kg"),
    ("Broccoli", "broccoli"),
    ("Tomatoes", "tomatoes 500g"),
    ("Baby spinach", "baby spinach 120g"),
    ("Eggs", "eggs 12 pack"),
    ("Chicken breast", "chicken breast 1kg"),
    ("Beef mince", "beef mince 500g"),
    ("Canned tuna", "tuna 95g"),
    ("Milk", "milk 2 litre"),
    ("Tasty cheese", "tasty cheese 500g"),
    ("Greek yoghurt", "greek yoghurt 1kg"),
    ("Wholemeal bread", "wholemeal bread loaf"),
    ("Rolled oats", "rolled oats 750g"),
    ("Pasta", "pasta 500g"),
    ("Rice", "rice 1kg"),
    ("Diced tomatoes", "diced tomatoes 400g"),
    ("Black beans", "black beans 400g"),
    ("Peanut butter", "peanut butter 375g"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)


def add_heading(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_numbered_step(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_after = Pt(3)
    n = p.add_run(f"{number}. ")
    n.font.name = "Arial"
    n.font.bold = True
    n.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Food Price Search Worksheet")
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    s = subtitle.add_run("Compare grocery prices on Coles and Woolworths websites")
    s.font.name = "Arial"
    s.font.size = Pt(11)

    info = doc.add_table(rows=2, cols=4)
    info.rows[0].cells[0].text = "Name"
    info.rows[0].cells[1].text = ""
    info.rows[0].cells[2].text = "Date"
    info.rows[0].cells[3].text = ""
    info.rows[1].cells[0].text = "Class"
    info.rows[1].cells[1].text = "RED-01 Pathway"
    info.rows[1].cells[2].text = "Trainer"
    info.rows[1].cells[3].text = "Simon"
    style_table(info)
    for cell in (info.rows[0].cells[0], info.rows[0].cells[2], info.rows[1].cells[0], info.rows[1].cells[2]):
        set_cell_shading(cell, "EDEDED")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    add_heading(doc, "Purpose")
    add_body(doc, "Use supermarket websites to search for food items, compare prices, and practise typing information into a worksheet.")

    safety = doc.add_table(rows=1, cols=1)
    safety.cell(0, 0).text = "Safety rule: Do not log in, create an account, add payment details, or buy anything."
    style_table(safety)
    set_cell_shading(safety.cell(0, 0), "FFF2CC")
    for run in safety.cell(0, 0).paragraphs[0].runs:
        run.font.bold = True

    add_heading(doc, "Websites")
    add_body(doc, "Coles: https://www.coles.com.au/")
    add_body(doc, "Woolworths: https://www.woolworths.com.au/shop/")

    add_heading(doc, "Steps")
    steps = [
        "Open the Coles or Woolworths website.",
        "Type the search words into the search box.",
        "Find the correct item and size.",
        "Write the price in the table.",
        "Compare the two prices.",
        "Write which store is cheaper, or write same price.",
    ]
    for i, step in enumerate(steps, start=1):
        add_numbered_step(doc, i, step)

    add_heading(doc, "Food Search Table")
    table = doc.add_table(rows=1, cols=6)
    headers = ["Food item", "Search words", "Coles price", "Woolworths price", "Cheaper store", "Notes"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
    for item, search in FOOD_ITEMS:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = search
        row[2].text = ""
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
    style_table(table)
    widths = [Cm(3.3), Cm(4.2), Cm(2.3), Cm(2.8), Cm(2.8), Cm(3.2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    add_heading(doc, "Reflection")
    reflection = [
        "Which item was the cheapest?",
        "Which item was more expensive than you expected?",
        "Did Coles or Woolworths have more cheaper prices today?",
        "What search words helped you find items faster?",
    ]
    for q in reflection:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(q)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.bold = True
        line = doc.add_paragraph("________________________________________________________________________________")
        line.paragraph_format.space_after = Pt(5)
        for rr in line.runs:
            rr.font.name = "Arial"
            rr.font.size = Pt(10)

    add_heading(doc, "Helpful words")
    vocab = doc.add_table(rows=1, cols=2)
    vocab.rows[0].cells[0].text = "Word"
    vocab.rows[0].cells[1].text = "Meaning"
    for cell in vocab.rows[0].cells:
        set_cell_shading(cell, "EDEDED")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
    for word, meaning in [
        ("brand", "the name of the product"),
        ("own brand", "the supermarket's cheaper brand"),
        ("in stock", "available to buy"),
        ("unit price", "price by weight or size to help compare value"),
        ("search", "type words to find an item online"),
    ]:
        row = vocab.add_row().cells
        row[0].text = word
        row[1].text = meaning
    style_table(vocab)
    vocab.rows[0].cells[0].width = Cm(4)
    vocab.rows[0].cells[1].width = Cm(11)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build()
