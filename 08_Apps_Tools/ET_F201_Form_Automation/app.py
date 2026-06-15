#!/usr/bin/env python3
from __future__ import annotations

import html
import json
import re
import urllib.parse
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "templates" / "ET F201.docx"
OUTPUT_DIR = ROOT / "Generated Forms"
DRAFTS_DIR = ROOT / "Saved Drafts"
HOST = "127.0.0.1"
PORT = 8765
BLUE_RGB = RGBColor(0x1F, 0x5F, 0xBF)
RATING_FONT = "Rethink Sans"
RATING_FONT_SIZE = Pt(9)

STREAMS = ["Initial", "Basic", "Advanced"]
LOADS = ["Full-time", "Part-time"]
BLOCK_LEVELS = [f"B{i}" for i in range(1, 8)]
LEVELS = [
    ("", "Not marked"),
    ("PLA", "PL A"),
    ("PLB", "PL B"),
    ("1", "Level 1"),
    ("2", "Level 2"),
    ("3", "Level 3"),
    ("4", "Level 4 (stored only)"),
]
OVERALL_RATINGS = [
    ("", "Not selected"),
    ("PLA", "PL A"),
    ("PLB", "PL B"),
    ("1", "Level 1"),
    ("2", "Level 2"),
    ("3", "Level 3"),
    ("4", "Level 4"),
]
INDICATORS = [
    ("01", "Learning .01"),
    ("02", "Learning .02"),
    ("03", "Reading .03"),
    ("04", "Reading .04"),
    ("05", "Writing .05"),
    ("06", "Writing .06"),
    ("07", "Oral Communication .07"),
    ("08", "Oral Communication .08"),
    ("09", "Numeracy .09"),
    ("10", "Numeracy .10"),
    ("11", "Numeracy .11"),
    ("12", "Digital Skills .12"),
    ("13", "Digital Skills .13"),
]
ROW_BY_INDICATOR = {
    "01": 2,
    "02": 3,
    "03": 5,
    "04": 6,
    "05": 8,
    "06": 9,
    "07": 11,
    "08": 12,
    "09": 14,
    "10": 15,
    "11": 16,
    "12": 18,
    "13": 19,
}
CODE_CELL = {"PLA": 0, "PLB": 2, "1": 4, "2": 6, "3": 8}
MARK_CELL = {"PLA": 1, "PLB": 3, "1": 5, "2": 7, "3": 8}
WORD_TEMPLATE_LEVELS = set(CODE_CELL)


def clean(value: str | None) -> str:
    return (value or "").strip()


def checkbox(label: str, selected: str) -> str:
    return f"{'☒' if label == selected else '☐'}  {label}"


def set_cell(cell, text: str, *, bold_label: bool = False, align_center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if bold_label and ":" in text:
        label, rest = text.split(":", 1)
        run = paragraph.add_run(label + ":")
        run.bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_blue_text(cell, text: str, *, size: int = 13, align_center: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.color.rgb = BLUE_RGB
    run.font.size = Pt(size)
    run.font.name = RATING_FONT
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_blue_dot(cell) -> None:
    set_cell_blue_text(cell, "•", size=14)


def selected_fields(post_body: bytes) -> dict[str, list[str]]:
    return urllib.parse.parse_qs(post_body.decode("utf-8"), keep_blank_values=True)


def one(form: dict[str, list[str]], key: str, default: str = "") -> str:
    return clean(form.get(key, [default])[0])


def all_values(form: dict[str, list[str]], key: str) -> list[str]:
    return [clean(value) for value in form.get(key, []) if clean(value)]


def flatten_form(form: dict[str, list[str]]) -> dict[str, str]:
    return {key: one(form, key) for key in form if not key.startswith("_")}


def form_from_flat(values: dict[str, str]) -> dict[str, list[str]]:
    return {key: [str(value)] for key, value in values.items()}


def participant_file_part(full_name: str) -> str:
    parts = [p for p in re.split(r"\s+", full_name.strip()) if p]
    if len(parts) >= 2:
        first = parts[0]
        surname = parts[-1]
        raw = f"{surname}-{first}"
    elif parts:
        raw = parts[0]
    else:
        raw = "Unknown-Participant"
    raw = re.sub(r"[^A-Za-z0-9-]+", "-", raw)
    return re.sub(r"-+", "-", raw).strip("-") or "Unknown-Participant"


def output_filename(form: dict[str, list[str]]) -> str:
    completed_date = one(form, "completion_date") or "Date-Unspecified"
    participant = participant_file_part(one(form, "participant_name"))
    return (
        f"{completed_date}_MTC_Redfern_RED-01_Pathway_"
        f"{participant}_ET-F201_Block-{one(form, 'block') or 'Unspecified'}.docx"
    )


def draft_filename(form: dict[str, list[str]]) -> str:
    completed_date = one(form, "completion_date") or "Date-Unspecified"
    participant = participant_file_part(one(form, "participant_name"))
    block = re.sub(r"[^A-Za-z0-9-]+", "-", one(form, "block") or "Unspecified").strip("-")
    return f"{completed_date}_{participant}_ET-F201_Block-{block or 'Unspecified'}.json"


def save_draft(form: dict[str, list[str]]) -> Path:
    DRAFTS_DIR.mkdir(exist_ok=True)
    path = DRAFTS_DIR / draft_filename(form)
    path.write_text(json.dumps(flatten_form(form), indent=2), encoding="utf-8")
    return path


def load_draft(name: str) -> dict[str, list[str]]:
    target = (DRAFTS_DIR / name).resolve()
    if target.parent != DRAFTS_DIR.resolve() or target.suffix != ".json":
        raise ValueError("Student record name is not valid.")
    data = json.loads(target.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Student record file is not valid.")
    return form_from_flat({str(k): str(v) for k, v in data.items()})


def saved_drafts() -> list[Path]:
    if not DRAFTS_DIR.exists():
        return []
    return sorted(DRAFTS_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)


def draft_values(path: Path) -> dict[str, str]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        return {}
    return {str(key): str(value) for key, value in data.items()}


def parse_outcome_code(code: str) -> tuple[str, str] | None:
    cleaned = clean(code).upper().replace(" ", "")
    match = re.fullmatch(r"(PLA|PLB|[1-4])\.(\d{1,2})", cleaned)
    if not match:
        return None
    level, indicator = match.groups()
    return level, indicator.zfill(2)


def outcome_codes_for_block(form: dict[str, list[str]], block: str) -> list[str]:
    raw = one(form, f"progressive_completed_{block}")
    return [part.strip() for part in re.split(r"[,;\n]+", raw) if part.strip()]


def ratings_to_circle(form: dict[str, list[str]]) -> list[str]:
    raw = one(form, "ratings_to_circle")
    return [part.strip() for part in raw.split(",") if part.strip()]


def block_number(block: str) -> int | None:
    match = re.fullmatch(r"B(\d+)", clean(block).upper())
    if not match:
        return None
    return int(match.group(1))


def blocks_through(selected_block: str) -> list[str]:
    selected_number = block_number(selected_block)
    if selected_number is None:
        return []
    return [block for block in BLOCK_LEVELS if (block_number(block) or 0) <= selected_number]


def next_block(block: str) -> str | None:
    selected_number = block_number(block)
    if selected_number is None:
        return None
    candidate = f"B{selected_number + 1}"
    return candidate if candidate in BLOCK_LEVELS else None


def block_hours(block: str) -> int | None:
    selected_number = block_number(block)
    if selected_number is None:
        return None
    return selected_number * 200


def prepare_next_block_form(form: dict[str, list[str]]) -> dict[str, list[str]]:
    current_block = one(form, "block")
    prepared_block = next_block(current_block)
    if not prepared_block:
        raise ValueError("This draft does not have a next block available.")
    prepared = {key: list(values) for key, values in form.items()}
    prepared["block"] = [prepared_block]
    hours = block_hours(prepared_block)
    if hours:
        prepared["block_note"] = [
            f"{prepared_block} reached after {hours} hours of attendance. This coversheet should reflect the Pre-Training Assessment baseline plus completed progressive assessment outcomes through {prepared_block}."
        ]
    return prepared


def cumulative_outcome_entries(form: dict[str, list[str]], selected_block: str) -> list[tuple[str, str]]:
    entries: list[tuple[str, str]] = []
    for block in blocks_through(selected_block):
        for outcome_code in outcome_codes_for_block(form, block):
            entries.append((block, outcome_code))
    return entries


def add_marker(cell, marker: str, *, prefix: str = "") -> None:
    existing = "\n".join(part for part in cell.text.splitlines() if part.strip())
    if marker in existing.splitlines():
        return
    if existing:
        set_cell(cell, f"{existing}\n{marker}", align_center=True)
    elif prefix:
        set_cell(cell, f"{prefix}\n{marker}", align_center=True)
    else:
        set_cell(cell, marker, align_center=True)


def mark_rating_cell(cell, rating_code: str, block: str, *, include_arrow: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dot_run = paragraph.add_run("●")
    dot_run.font.color.rgb = BLUE_RGB
    dot_run.font.size = Pt(10)
    space_run = paragraph.add_run(" ")
    space_run.font.name = RATING_FONT
    space_run.font.size = RATING_FONT_SIZE
    rating_run = paragraph.add_run(rating_code)
    rating_run.font.name = RATING_FONT
    rating_run.font.size = RATING_FONT_SIZE
    if include_arrow:
        arrow_run = paragraph.add_run(f" {block or 'B__'} →")
        arrow_run.font.color.rgb = BLUE_RGB
        arrow_run.font.size = RATING_FONT_SIZE
        arrow_run.font.name = RATING_FONT


def mark_arrow_cell(cell, block: str) -> None:
    set_cell_blue_text(cell, f"{block or 'B__'}\n→", size=9)


def mark_outcome_on_grid(table, outcome_code: str, block: str) -> str | None:
    parsed = parse_outcome_code(outcome_code)
    if not parsed:
        return f"{outcome_code} is not a valid outcome code."
    level, indicator = parsed
    if level not in WORD_TEMPLATE_LEVELS:
        return f"{outcome_code} is saved but cannot be shown in this ET F201 template."
    if indicator not in ROW_BY_INDICATOR:
        return f"{outcome_code} does not match an ET F201 indicator."
    row = ROW_BY_INDICATOR[indicator]
    rating_code = f"{level}.{indicator}"
    mark_rating_cell(table.cell(row, CODE_CELL[level]), rating_code, block, include_arrow=(level == "3"))
    if level != "3":
        mark_arrow_cell(table.cell(row, MARK_CELL[level]), block)
    return None


def mark_rating_on_grid(table, rating_code: str) -> str | None:
    parsed = parse_outcome_code(rating_code)
    if not parsed:
        return f"{rating_code} is not a valid ACSF rating code."
    level, indicator = parsed
    if level not in WORD_TEMPLATE_LEVELS:
        return f"{rating_code} is saved but cannot be shown in this ET F201 template."
    if indicator not in ROW_BY_INDICATOR:
        return f"{rating_code} could not be found in the ET F201 rating grid."
    row = ROW_BY_INDICATOR[indicator]
    mark_rating_cell(table.cell(row, CODE_CELL[level]), f"{level}.{indicator}", "", include_arrow=False)
    return None


def progressive_outcome_issues(form: dict[str, list[str]]) -> list[str]:
    issues = []
    for block in BLOCK_LEVELS:
        for outcome_code in outcome_codes_for_block(form, block):
            parsed = parse_outcome_code(outcome_code)
            if not parsed:
                issues.append(f"{block}: {outcome_code} is not a valid outcome code.")
                continue
            level, indicator = parsed
            if indicator not in ROW_BY_INDICATOR:
                issues.append(f"{block}: {outcome_code} does not match an ET F201 indicator.")
            if level not in WORD_TEMPLATE_LEVELS:
                issues.append(f"{block}: {outcome_code} is saved but cannot be shown in this ET F201 template.")
    return issues


def ratings_to_circle_issues(form: dict[str, list[str]]) -> list[str]:
    issues = []
    for rating_code in ratings_to_circle(form):
        parsed = parse_outcome_code(rating_code)
        if not parsed:
            issues.append(f"{rating_code} is not a valid ACSF rating code.")
            continue
        level, indicator = parsed
        if indicator not in ROW_BY_INDICATOR:
            issues.append(f"{rating_code} could not be found in the ET F201 rating grid.")
        if level not in WORD_TEMPLATE_LEVELS:
            issues.append(f"{rating_code} is saved but cannot be shown in this ET F201 template.")
    return issues


def overall_rating_value(form: dict[str, list[str]]) -> str:
    fixed_rating = one(form, "overall_rating")
    if fixed_rating:
        return fixed_rating
    for block in BLOCK_LEVELS:
        block_rating = one(form, f"overall_rating_{block}")
        if block_rating:
            return block_rating
    return ""


def missing_field_labels(form: dict[str, list[str]]) -> list[str]:
    checks = [
        ("block", "Completed block level"),
        ("participant_id", "JSID / Participant ID"),
        ("crn", "CRN"),
    ]
    missing = [label for key, label in checks if not one(form, key)]
    completed_block = one(form, "block")
    if completed_block and not one(form, f"progressive_completed_{completed_block}"):
        missing.append(f"Progressive assessment completed outcomes for {completed_block}")
    if not any(one(form, f"level_{indicator}") for indicator, _label in INDICATORS):
        missing.append("PTA starting scores")
    return missing


def unsupported_template_levels(form: dict[str, list[str]]) -> list[str]:
    unsupported = []
    for indicator, label in INDICATORS:
        level = one(form, f"level_{indicator}")
        if level and level not in WORD_TEMPLATE_LEVELS:
            unsupported.append(f"{label}: Level {level}")
    return unsupported


def review_message(form: dict[str, list[str]]) -> str:
    participant = one(form, "participant_name") or "this student"
    missing = missing_field_labels(form)
    unsupported = unsupported_template_levels(form)
    outcome_issues = progressive_outcome_issues(form)
    circle_issues = ratings_to_circle_issues(form)
    if not missing and not unsupported and not outcome_issues and not circle_issues:
        return (
            f'<div class="message">No missing core fields found for {html.escape(participant)}. '
            "Student record has the core fields needed for the coversheet draft. Review the Word draft before using it in any official system.</div>"
        )
    items = "".join(f"<li>{html.escape(item)}</li>" for item in missing)
    if unsupported:
        items += "<li>Current ET F201 Word template has no Level 4 column for: "
        items += html.escape(", ".join(unsupported))
        items += "</li>"
    if outcome_issues:
        items += "".join(f"<li>{html.escape(item)}</li>" for item in outcome_issues)
    if circle_issues:
        items += "".join(f"<li>{html.escape(item)}</li>" for item in circle_issues)
    return (
        f'<div class="message">Still to confirm for {html.escape(participant)}:'
        f"<ul>{items}</ul></div>"
    )


def fill_document(form: dict[str, list[str]]) -> tuple[Path, list[str]]:
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document(TEMPLATE)

    stream = one(form, "tuition_stream")
    training_load = one(form, "training_load")
    block = one(form, "block")
    trainer = one(form, "trainer_name") or "Simon Judge"

    t1 = doc.tables[0]
    for i, label in enumerate(STREAMS, start=1):
        set_cell(t1.cell(0, i), checkbox(label, stream), align_center=True)
    set_cell(t1.cell(1, 1), block or "______________", align_center=True)

    t2 = doc.tables[1]
    set_cell(t2.cell(0, 0), f"Participant's Full Name: {one(form, 'participant_name')}", bold_label=True)
    set_cell(t2.cell(1, 0), f"Participant ID: {one(form, 'participant_id')}", bold_label=True)
    set_cell(t2.cell(1, 3), f"CRN: {one(form, 'crn')}", bold_label=True)
    set_cell(t2.cell(2, 1), checkbox("Full-time", training_load), align_center=True)
    set_cell(t2.cell(2, 2), checkbox("Part-time", training_load), align_center=True)
    set_cell(t2.cell(2, 3), f"Trainer's Name: {trainer}", bold_label=True)

    t3 = doc.tables[2]
    template_notes: list[str] = []
    for indicator, label in INDICATORS:
        level = one(form, f"level_{indicator}")
        if not level:
            continue
        if level not in WORD_TEMPLATE_LEVELS:
            template_notes.append(f"{label} Level {level} is saved in the student record but not shown in this ET F201 template.")
            continue
        row = ROW_BY_INDICATOR[indicator]
        code_cell = CODE_CELL[level]
        code = f"{level}.{indicator}" if level in {"1", "2", "3"} else f"{level}.{indicator}"
        mark_rating_cell(t3.cell(row, code_cell), code, block, include_arrow=False)

    completed_entries = cumulative_outcome_entries(form, block) if block else []
    for completed_block, outcome_code in completed_entries:
        note = mark_outcome_on_grid(t3, outcome_code, completed_block)
        if note:
            template_notes.append(note)

    circle_warnings: list[str] = []
    selected_ratings_to_circle = ratings_to_circle(form)
    for rating_code in selected_ratings_to_circle:
        note = mark_rating_on_grid(t3, rating_code)
        if note:
            circle_warnings.append(note)
            template_notes.append(note)

    increase_summary = one(form, "increase_summary")

    t4 = doc.tables[3]
    set_cell(
        t4.cell(1, 0),
        "ACSF/DLSF Increases:\n" + increase_summary,
        bold_label=True,
    )
    set_cell(
        t4.cell(1, 2),
        "CGEA / FSK / EAL Unit(s):\n" + one(form, "units"),
        bold_label=True,
    )
    t5 = doc.tables[4]
    set_cell(t5.cell(1, 0), f"{'☒' if one(form, 'pa_ctp_checked') == 'yes' else '☐'}  PA / CTP Checked")
    set_cell(t5.cell(1, 1), f"{'☒' if one(form, 'soc_checked') == 'yes' else '☐'}  SOC Checked")
    set_cell(t5.cell(3, 0), f"Print Name: {one(form, 'senior_print_name')}", bold_label=True)
    set_cell(t5.cell(3, 1), f"Signature: {one(form, 'senior_signature')}", bold_label=True)

    output_path = OUTPUT_DIR / output_filename(form)
    doc.save(output_path)
    return output_path, circle_warnings


def field(value: str) -> str:
    return html.escape(value, quote=True)


def selected_attr(current: str, option: str) -> str:
    return " selected" if current == option else ""


def checked_attr(current: str, option: str = "yes") -> str:
    return " checked" if current == option else ""


def value(form: dict[str, list[str]], key: str, default: str = "") -> str:
    return field(one(form, key, default))


def textarea_value(form: dict[str, list[str]], key: str) -> str:
    return html.escape(one(form, key))


def raw_score_label(level: str, indicator: str) -> str:
    return f"{level}.{indicator}"


def drafts_panel() -> str:
    drafts = saved_drafts()
    if not drafts:
        return '<p class="muted">No saved student records yet.</p>'
    rows = []
    for draft in drafts[:20]:
        values = draft_values(draft)
        load_link = f"/draft?name={urllib.parse.quote(draft.name)}"
        next_link = f"/next-block?name={urllib.parse.quote(draft.name)}"
        participant = values.get("participant_name") or draft.stem.replace("_", " ")
        block = values.get("block") or "Not selected"
        next_action = (
            f'<a href="{next_link}">Prepare next block</a>'
            if next_block(block)
            else '<span class="muted">No next block</span>'
        )
        rows.append(
            "<tr>"
            f"<th scope=\"row\">{html.escape(participant)}</th>"
            f"<td>{html.escape(block)}</td>"
            f"<td><a href=\"{load_link}\">Load</a></td>"
            f"<td>{next_action}</td>"
            "</tr>"
        )
    return (
        "<table class=\"draft-list\">"
        "<thead><tr><th>Student</th><th>Block</th><th>Record</th><th>Next 200-hour block</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody>"
        "</table>"
    )


def render_form(message: str = "", values: dict[str, list[str]] | None = None) -> bytes:
    values = values or {}
    selected_date = one(values, "completion_date")
    indicator_rows = []
    for indicator, label in INDICATORS:
        current_level = one(values, f"level_{indicator}")
        rating_cells = []
        for option_value, text in LEVELS:
            if not option_value:
                continue
            score_label = raw_score_label(option_value, indicator)
            rating_cells.append(
                f"""
                <td>
                  <label class="rating-choice">
                    <input type="radio" name="level_{indicator}" value="{field(option_value)}"{checked_attr(current_level, option_value)}>
                    <span>{field(score_label)}</span>
                  </label>
                </td>
                """
            )
        indicator_rows.append(
            f"""
            <tr>
              {''.join(rating_cells)}
            </tr>
            """
        )

    current_stream = one(values, "tuition_stream")
    stream_options = "\n".join(
        f'<label class="choice"><input type="radio" name="tuition_stream" value="{s}"{checked_attr(current_stream, s)}> {s}</label>'
        for s in STREAMS
    )
    current_load = one(values, "training_load")
    load_options = "\n".join(
        f'<label class="choice"><input type="radio" name="training_load" value="{s}"{checked_attr(current_load, s)}> {s}</label>'
        for s in LOADS
    )
    current_block = one(values, "block")
    block_options = "\n".join(
        f'<option value="{field(block)}"{selected_attr(current_block, block)}>{field(block)}</option>'
        for block in BLOCK_LEVELS
    )
    progressive_completion_rows = []
    for block in BLOCK_LEVELS:
        progressive_completion_rows.append(
            f"""
            <tr>
              <th scope="row">{field(block)}</th>
              <td><input type="text" name="progressive_completed_{block}" value="{value(values, f'progressive_completed_{block}')}" placeholder="Example cells to mark: PLB.13, 1.01"></td>
              <td><input type="text" name="progressive_note_{block}" value="{value(values, f'progressive_note_{block}')}" placeholder="Optional source note"></td>
            </tr>
            """
        )
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>ET F201 Form Automation</title>
  <style>
    :root {{
      --yellow: #ffd200;
      --black: #111111;
      --line: #d8d8d8;
      --soft: #f6f6f2;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      font-family: Arial, Helvetica, sans-serif;
      color: var(--black);
      background: #ffffff;
    }}
    header {{
      background: var(--yellow);
      border-bottom: 4px solid var(--black);
      padding: 22px 28px;
    }}
    header h1 {{
      margin: 0;
      font-size: 30px;
      letter-spacing: 0;
    }}
    main {{
      width: min(1120px, calc(100% - 32px));
      margin: 24px auto 48px;
    }}
    form {{
      display: grid;
      gap: 18px;
    }}
    section {{
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 18px;
      background: #fff;
    }}
    h2 {{
      margin: 0 0 14px;
      font-size: 19px;
    }}
    .grid {{
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 14px;
    }}
    label span {{
      display: block;
      font-weight: 700;
      margin-bottom: 6px;
    }}
    input[type="text"], input[type="date"], select, textarea {{
      width: 100%;
      min-height: 42px;
      border: 1px solid #b8b8b8;
      border-radius: 6px;
      padding: 8px 10px;
      font: inherit;
      background: #fff;
    }}
    textarea {{ min-height: 92px; resize: vertical; }}
    .choices {{
      display: flex;
      gap: 14px;
      flex-wrap: wrap;
      align-items: center;
      min-height: 42px;
    }}
    .choice, .tick {{
      display: inline-flex;
      align-items: center;
      gap: 7px;
      font-weight: 400;
    }}
    .rating-choice {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 6px;
      width: 100%;
      min-height: 34px;
      font-weight: 400;
      white-space: nowrap;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
    }}
    th, td {{
      border-bottom: 1px solid var(--line);
      padding: 9px 8px;
      text-align: left;
      vertical-align: middle;
    }}
    th {{ width: 38%; }}
    td:nth-child(2) {{ width: 32%; }}
    td:nth-child(3) {{ width: 30%; }}
    .pta-grid td {{
      width: 16.66%;
      text-align: center;
    }}
    .note {{
      background: var(--soft);
      border-left: 6px solid var(--yellow);
      padding: 12px 14px;
      margin: 0 0 14px;
    }}
    .actions {{
      display: flex;
      gap: 12px;
      align-items: center;
      flex-wrap: wrap;
    }}
    button, .button {{
      appearance: none;
      border: 2px solid var(--black);
      background: var(--black);
      color: #fff;
      border-radius: 6px;
      padding: 11px 16px;
      font-weight: 700;
      text-decoration: none;
      cursor: pointer;
      font-size: 16px;
    }}
    .button.secondary {{
      background: #fff;
      color: var(--black);
    }}
    .message {{
      border: 2px solid var(--black);
      background: var(--yellow);
      padding: 14px 16px;
      border-radius: 8px;
      font-weight: 700;
    }}
    .muted {{
      color: #555;
      margin: 0;
    }}
    .drafts ul {{
      margin: 0;
      padding-left: 20px;
      display: grid;
      gap: 6px;
    }}
    .drafts a {{
      color: var(--black);
      font-weight: 700;
    }}
    .draft-list th {{
      width: auto;
    }}
    .draft-list td:nth-child(2),
    .draft-list td:nth-child(3),
    .draft-list td:nth-child(4) {{
      width: auto;
    }}
    @media (max-width: 780px) {{
      .grid {{ grid-template-columns: 1fr; }}
      table, tbody, tr, th, td {{ display: block; width: 100%; }}
      th, td {{ border-bottom: 0; padding: 7px 0; }}
      tr {{ border-bottom: 1px solid var(--line); padding: 8px 0; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>ET F201 Form Automation</h1>
  </header>
  <main>
    {message}
    <form method="post" action="/generate">
      <section class="drafts">
        <h2>Saved Student Records</h2>
        <p class="note">Each student you set up is saved as a local student record. Reopen the record later to create the next coversheet without re-entering the fixed details or PTA baseline.</p>
        {drafts_panel()}
      </section>

      <section>
        <h2>Student Setup</h2>
        <p class="note">Enter the core student details used to prepare the SEE Assessment Coversheet. Leave unconfirmed fields blank and review the Word draft before uploading or submitting it.</p>
        <div class="grid">
          <label><span>Student full name</span><input type="text" name="participant_name" value="{value(values, 'participant_name')}"></label>
          <label><span>JSID / Participant ID</span><input type="text" name="participant_id" value="{value(values, 'participant_id')}"></label>
          <label><span>CRN</span><input type="text" name="crn" value="{value(values, 'crn')}"></label>
          <label><span>Completed block level</span><select name="block"><option value="">Not selected</option>{block_options}</select></label>
          <label><span>Trainer's name</span><input type="text" name="trainer_name" value="{value(values, 'trainer_name', 'Simon Judge')}"></label>
          <label><span>Date</span><input type="date" name="completion_date" value="{field(selected_date)}"></label>
          <label><span>Block note</span><textarea name="block_note" placeholder="Saved in the student record only. Example: B1 reached after 200 hours of attendance.">{textarea_value(values, 'block_note')}</textarea></label>
        </div>
      </section>

      <section>
        <h2>Stream and Training Load</h2>
        <div class="grid">
          <div><label><span>Tuition stream</span></label><div class="choices">{stream_options}</div></div>
          <div><label><span>Training load</span></label><div class="choices">{load_options}</div></div>
          <label><span>Stream calculation note</span><textarea name="stream_calculation_note" placeholder="Saved in the student record only. Example: Basic stream calculated using rule of 7.">{textarea_value(values, 'stream_calculation_note')}</textarea></label>
        </div>
      </section>

      <section>
        <h2>Student ACSF Record</h2>
        <p class="note">Enter the Pre-Training Assessment once as the student's baseline. For each later block, add only the new progressive assessment outcomes. When a block is selected, the Word draft shows the PTA baseline plus all completed block outcomes up to that point. Level 4 can be stored in the app, but the current ET F201 Word template only shows columns up to Level 3.</p>
        <h3>PTA Starting Scores / One-Time Coversheet Baseline</h3>
        <p class="note">Tick the raw PTA score for each indicator, such as PLA.01, PLB.01, 1.01, or 2.01. The score already identifies the level. These selections are saved as the student's baseline and appear on each coversheet created from the student record.</p>
        <table class="pta-grid">
          <thead><tr><th>PL A score</th><th>PL B score</th><th>Level 1 score</th><th>Level 2 score</th><th>Level 3 score</th><th>Level 4 score</th></tr></thead>
          <tbody>{''.join(indicator_rows)}</tbody>
        </table>
        <h3>What ACSF ratings do you want circled?</h3>
        <p class="note">Type one or more ACSF ratings separated by commas. The Word draft will visually mark each matching rating cell.</p>
        <label>
          <span>What ACSF ratings do you want circled?</span>
          <textarea name="ratings_to_circle" placeholder="PLB.13, 1.01, 1.03, 2.07">{textarea_value(values, 'ratings_to_circle')}</textarea>
        </label>
        <h3>Circle / Mark ACSF Ratings Achieved By Block</h3>
        <p class="note">For each block, enter the exact ACSF rating cells that need to be marked on the coversheet. Enter codes separated by commas, for example: PLB.13, 1.01. The selected block is cumulative, so B2 includes B1 and B2, B3 includes B1 to B3, and so on.</p>
        <table>
          <thead><tr><th>Block</th><th>ACSF rating cells to mark</th><th>Source note</th></tr></thead>
          <tbody>{''.join(progressive_completion_rows)}</tbody>
        </table>
      </section>

      <section>
        <h2>Outcomes and Checks</h2>
        <div class="grid">
          <label><span>ACSF/DLSF increases summary</span><textarea name="increase_summary" placeholder="Enter ACSF/DLSF increases manually.">{textarea_value(values, 'increase_summary')}</textarea></label>
          <label><span>CGEA / FSK / EAL units</span><textarea name="units">{textarea_value(values, 'units')}</textarea></label>
        </div>
      </section>

      <div class="actions">
        <button type="submit" formaction="/save-record">Save Student Record</button>
        <button type="submit" formaction="/review">Review Missing Fields</button>
        <button type="submit">Create Draft Word Document</button>
        <a class="button secondary" href="/folder">Open Generated Forms Folder</a>
      </div>
    </form>
  </main>
  <script>
    document.querySelectorAll('.pta-grid input[type="radio"]').forEach((radio) => {{
      const rememberState = () => {{
        radio.dataset.wasChecked = radio.checked ? 'yes' : 'no';
      }};
      radio.addEventListener('pointerdown', rememberState);
      radio.addEventListener('mousedown', rememberState);
      radio.addEventListener('click', (event) => {{
        if (radio.dataset.wasChecked === 'yes') {{
          event.preventDefault();
          window.setTimeout(() => {{
            radio.checked = false;
            radio.dispatchEvent(new Event('change', {{ bubbles: true }}));
          }}, 0);
        }}
        radio.dataset.wasChecked = radio.checked ? 'yes' : 'no';
      }});
    }});
  </script>
</body>
</html>"""
    return html_doc.encode("utf-8")


class Handler(BaseHTTPRequestHandler):
    def do_GET(self) -> None:
        if self.path == "/" or self.path.startswith("/?"):
            self.send_html(render_form())
            return
        if self.path.startswith("/draft?"):
            query = urllib.parse.parse_qs(urllib.parse.urlsplit(self.path).query)
            name = one(query, "name")
            try:
                values = load_draft(name)
            except Exception as exc:
                message = f'<p class="message">Could not load student record: {html.escape(str(exc))}</p>'
                self.send_html(render_form(message))
                return
            message = f'<p class="message">Loaded student record: {html.escape(name)}</p>'
            self.send_html(render_form(message, values))
            return
        if self.path.startswith("/next-block?"):
            query = urllib.parse.parse_qs(urllib.parse.urlsplit(self.path).query)
            name = one(query, "name")
            try:
                values = prepare_next_block_form(load_draft(name))
            except Exception as exc:
                message = f'<p class="message">Could not prepare the next block: {html.escape(str(exc))}</p>'
                self.send_html(render_form(message))
                return
            message = (
                f'<p class="message">Prepared the next block from student record {html.escape(name)}. '
                "Review the new block details, add the completed outcomes for that block, then save the student record.</p>"
            )
            self.send_html(render_form(message, values))
            return
        if self.path == "/folder":
            import subprocess

            OUTPUT_DIR.mkdir(exist_ok=True)
            subprocess.run(["open", str(OUTPUT_DIR)], check=False)
            self.send_html(render_form('<p class="message">Generated Forms folder opened in Finder.</p>'))
            return
        if self.path.startswith("/download/"):
            filename = urllib.parse.unquote(self.path.removeprefix("/download/"))
            target = OUTPUT_DIR / filename
            if target.exists() and target.parent == OUTPUT_DIR:
                data = target.read_bytes()
                self.send_response(200)
                self.send_header(
                    "Content-Type",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
                self.send_header("Content-Length", str(len(data)))
                self.end_headers()
                self.wfile.write(data)
                return
        self.send_error(404)

    def do_POST(self) -> None:
        if self.path not in {"/generate", "/save-draft", "/save-record", "/review"}:
            self.send_error(404)
            return
        length = int(self.headers.get("Content-Length", "0"))
        form = selected_fields(self.rfile.read(length))
        if self.path == "/review":
            self.send_html(render_form(review_message(form), form))
            return
        if self.path in {"/save-draft", "/save-record"}:
            try:
                draft = save_draft(form)
            except Exception as exc:
                self.send_html(render_form(f'<p class="message">Could not save the student record: {html.escape(str(exc))}</p>', form))
                return
            message = f'<p class="message">Saved student record: {html.escape(draft.name)}</p>'
            self.send_html(render_form(message, form))
            return
        try:
            output, warnings = fill_document(form)
        except Exception as exc:
            self.send_html(render_form(f'<p class="message">Could not create the document: {html.escape(str(exc))}</p>', form))
            return
        link = f"/download/{urllib.parse.quote(output.name)}"
        warning_html = ""
        if warnings:
            warning_items = "".join(f"<li>{html.escape(warning)}</li>" for warning in warnings)
            warning_html = f"<br>Warning: the following ratings could not be marked:<ul>{warning_items}</ul>"
        message = (
            '<p class="message">Word document created: '
            f'<a href="{link}">{html.escape(output.name)}</a>{warning_html}</p>'
        )
        self.send_html(render_form(message, form))

    def send_html(self, body: bytes) -> None:
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format: str, *args) -> None:
        return


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(exist_ok=True)
    DRAFTS_DIR.mkdir(exist_ok=True)
    print(f"ET F201 form automation is running at http://{HOST}:{PORT}")
    print("Close this window to stop the form.")
    ThreadingHTTPServer((HOST, PORT), Handler).serve_forever()
